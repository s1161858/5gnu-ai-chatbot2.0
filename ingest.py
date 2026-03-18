#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ingest.py - 5GNU RAG 知识库构建脚本
功能：读取本地 PDF/Word 文件 + 爬取公司网站，向量化后存入 ChromaDB
"""

import os
import logging
from pathlib import Path
from typing import List

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

# ──────────────────────────────────────────────
# 基础配置
# ──────────────────────────────────────────────
load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

DATA_DIR         = Path("./data")
VECTOR_STORE_DIR = "./vector_store"
COLLECTION_NAME  = "5gnu_knowledge"
CHUNK_SIZE       = 500
CHUNK_OVERLAP    = 100
EMBEDDING_MODEL  = "BAAI/bge-large-zh-v1.5"
COMPANY_URLS: List[str] = [
    "https://ltexpo2023.5gnumultimedia.com/",
]


# ──────────────────────────────────────────────
# 1. PDF 文档加载
# ──────────────────────────────────────────────
def load_pdf(file_path: Path) -> List[Document]:
    docs = []
    try:
        logger.info(f"📄 正在处理 PDF：{file_path.name}")
        pdf = fitz.open(str(file_path))
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(file_path), "page": page_num, "type": "pdf"}
                ))
        pdf.close()
        logger.info(f"   ✅ 成功提取 {len(docs)} 页内容")
    except Exception as e:
        logger.error(f"   ❌ PDF 解析失败 [{file_path.name}]：{e}")
    return docs


# ──────────────────────────────────────────────
# 2. Word 文档加载
# ──────────────────────────────────────────────
def load_docx(file_path: Path) -> List[Document]:
    docs = []
    try:
        logger.info(f"📝 正在处理 Word：{file_path.name}")
        docx = DocxDocument(str(file_path))
        full_text = "\n".join(
            para.text.strip()
            for para in docx.paragraphs
            if para.text.strip()
        )
        if full_text:
            docs.append(Document(
                page_content=full_text,
                metadata={"source": str(file_path), "type": "docx"}
            ))
        logger.info(f"   ✅ 成功提取 {len(full_text)} 字符")
    except Exception as e:
        logger.error(f"   ❌ Word 解析失败 [{file_path.name}]：{e}")
    return docs


# ──────────────────────────────────────────────
# 3. 遍历 /data 目录
# ──────────────────────────────────────────────
def load_all_documents() -> List[Document]:
    all_docs = []
    if not DATA_DIR.exists():
        logger.warning(f"⚠️  数据目录不存在：{DATA_DIR}")
        return all_docs

    for file_path in DATA_DIR.iterdir():
        if file_path.suffix.lower() == ".pdf":
            all_docs.extend(load_pdf(file_path))
        elif file_path.suffix.lower() == ".docx":
            all_docs.extend(load_docx(file_path))

    logger.info(f"📚 本地文档加载完成，共 {len(all_docs)} 个文档块")
    return all_docs


# ──────────────────────────────────────────────
# 4. 网页爬取
# ──────────────────────────────────────────────
def scrape_url(url: str) -> List[Document]:
    docs = []
    try:
        logger.info(f"🌐 正在爬取网页：{url}")
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            )
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        response.encoding = response.apparent_encoding

        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "iframe", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        lines = [line for line in text.splitlines() if line.strip()]
        clean_text = "\n".join(lines)

        if clean_text:
            docs.append(Document(
                page_content=clean_text,
                metadata={"source": url, "type": "webpage"}
            ))
            logger.info(f"   ✅ 爬取网页成功，提取 {len(clean_text)} 字符")
    except Exception as e:
        logger.error(f"   ❌ 网页爬取失败 [{url}]：{e}")
    return docs


def scrape_all_urls(url_list: List[str]) -> List[Document]:
    all_web_docs = []
    for url in url_list:
        all_web_docs.extend(scrape_url(url))
    return all_web_docs


# ──────────────────────────────────────────────
# 5. 文本切片
# ──────────────────────────────────────────────
def split_documents(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    logger.info(f"✂️  文本切片完成，共生成 {len(chunks)} 个切片")
    return chunks


# ──────────────────────────────────────────────
# 6. 向量化与存储
# ──────────────────────────────────────────────
def build_vector_store(chunks: List[Document]) -> None:
    if not chunks:
        logger.error("❌ 没有可向量化的切片")
        return

    logger.info(f"🤖 加载 Embedding 模型：{EMBEDDING_MODEL}")
    embedding_func = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    logger.info(f"💾 写入 ChromaDB → {VECTOR_STORE_DIR}")
    Chroma.from_documents(
        documents=chunks,
        embedding=embedding_func,
        collection_name=COLLECTION_NAME,
        persist_directory=VECTOR_STORE_DIR,
    )
    logger.info(f"✅ 向量库构建完成！共写入 {len(chunks)} 条记录")


# ──────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────
def main():
    logger.info("=" * 55)
    logger.info("  🚀 5GNU RAG 知识库构建开始")
    logger.info("=" * 55)

    local_docs = load_all_documents()
    web_docs   = scrape_all_urls(COMPANY_URLS)
    all_docs   = local_docs + web_docs

    if not all_docs:
        logger.error("❌ 未获取到任何文档")
        return

    chunks = split_documents(all_docs)
    build_vector_store(chunks)

    logger.info("=" * 55)
    logger.info("  🎉 知识库构建完成！")
    logger.info("=" * 55)


if __name__ == "__main__":
    main()
